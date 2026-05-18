"""
FastAPI Backend for Genie Local Chatbot
"""

import sys
import io
import re
if sys.platform == "win32":
    try:
        sys.stdout.reconfigure(encoding="utf-8", errors="replace")
        sys.stderr.reconfigure(encoding="utf-8", errors="replace")
    except AttributeError:
        if hasattr(sys.stdout, "buffer"):
            sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace", line_buffering=True)
        if hasattr(sys.stderr, "buffer"):
            sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding="utf-8", errors="replace", line_buffering=True)

import json
import time
import os
import re
import threading
from typing import Optional
from pathlib import Path
import uuid
from contextlib import asynccontextmanager
import zipfile
import xml.etree.ElementTree as ET

import traceback

from fastapi import FastAPI, HTTPException, UploadFile, File, Form
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse, JSONResponse
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel, Field

try:
    from PIL import Image
except ImportError:
    Image = None

try:
    import PyPDF2
except ImportError:
    PyPDF2 = None

try:
    import docx
except ImportError:
    docx = None

try:
    import olefile
except ImportError:
    olefile = None

from backend.model_manager import model_manager
from backend.chat_engine import chat_engine, ChatSession


def sanitize_for_console(text: str) -> str:
    """Remove or replace characters that Windows console can't handle."""
    import re
    
    emoji_map = {
        '\U0001f4c1': '[FILE]',   # 📁
        '\U0001f4c4': '[DOC]',    # 📄
        '\U0001f4be': '[SAVE]',   # 💾
        '\U0001f4f7': '[IMG]',    # 📷
        '\U0001f50d': '[SEARCH]', # 🔍
        '\U0001f4dd': '[NOTE]',   # 📝
    }
    
    for emoji, replacement in emoji_map.items():
        text = text.replace(emoji, replacement)
    
    # Remove any other non-BMP characters (Windows console limitation)
    text = re.sub(r'[^\u0000-\uFFFF]', '?', text)
    
    return text


# ── Lifespan: auto-load model ────────────────────────────────────
@asynccontextmanager
async def lifespan(app: FastAPI):
    model_manager.auto_load()
    yield


app = FastAPI(
    title="Genie Chatbot",
    version="1.0.0",
    lifespan=lifespan,
)

# CORS - MUST be before routes
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

@app.middleware("http")
async def log_requests(request, call_next):
    start = time.time()
    print(f"[IN] {request.method} {request.url}")
    try:
        response = await call_next(request)
        duration = time.time() - start
        print(f"[OUT] {response.status_code} {request.method} {request.url} ({duration:.2f}s)")
        return response
    except Exception as exc:
        duration = time.time() - start
        print(f"[ERROR] {request.method} {request.url} ({duration:.2f}s): {exc}")
        traceback.print_exc()
        raise


@app.exception_handler(Exception)
async def global_exception_handler(request, exc):
    print(f"[ERROR] UNHANDLED ERROR: {exc}")
    traceback.print_exc()
    return JSONResponse(
        status_code=500,
        content={"detail": str(exc)}
    )


# ── Serve frontend ───────────────────────────────────────────────
frontend_path = Path(__file__).parent.parent / "frontend"
if frontend_path.exists():
    app.mount("/static", StaticFiles(directory=str(frontend_path)), name="static")

    @app.get("/app", include_in_schema=False)
    async def serve_app():
        from fastapi.responses import FileResponse
        return FileResponse(str(frontend_path / "index.html"))

# ── Serve uploads ────────────────────────────────────────────────
uploads_path = Path(__file__).parent / "uploads"
uploads_path.mkdir(exist_ok=True)
app.mount("/uploads", StaticFiles(directory=str(uploads_path)), name="uploads")

# ── Request models ───────────────────────────────────────────────
class ChatRequest(BaseModel):
    message: str = Field(..., min_length=1, max_length=10000)
    session_id: Optional[str] = None
    temperature: Optional[float] = Field(1.0, ge=0.0, le=2.0)
    max_tokens: Optional[int] = Field(2048, ge=1, le=8192)
    enable_thinking: Optional[bool] = False
    system_prompt: Optional[str] = None


class SettingsUpdate(BaseModel):
    temperature: Optional[float] = None
    top_p: Optional[float] = None
    top_k: Optional[int] = None
    max_tokens: Optional[int] = None


# ── Status ───────────────────────────────────────────────────────
@app.get("/")
async def root():
    return {"status": "running", "model": model_manager.get_status()}


@app.get("/api/status")
async def get_status():
    return {
        "model": model_manager.get_status(),
        "active_sessions": len(chat_engine.sessions),
        "uptime": time.time(),
    }


# ── Chat ─────────────────────────────────────────────────────────
@app.post("/api/chat")
async def chat(request: ChatRequest):
    if not model_manager.model_loaded:
        detail = "Model is still loading..." if model_manager.loading else model_manager.load_error
        raise HTTPException(status_code=503, detail=detail)

    session = _get_or_create_session(request.session_id, request.system_prompt)
    extract_user_facts(session, request.message)
    session.add_message("user", request.message)
    history = session.get_history_for_model()

    result = model_manager.generate_sync(
        messages=history,
        temperature=1.0,
        max_tokens=8192,
    )
    if "error" in result:
        raise HTTPException(status_code=500, detail=result["error"])

    session.add_message("assistant", result["response"])

    # ── TRIGGER BACKGROUND COMPRESSION (ALL PAST MESSAGES) ───
    _trigger_compression_async(session)

    return {
        "response": result["response"],
        "session_id": session.session_id,
        "usage": result.get("usage", {}),
        "should_suggest_new_chat": session.should_suggest_new_chat(),
    }


@app.post("/api/chat/stream")
async def chat_stream(request: ChatRequest):
    if not model_manager.model_loaded:
        msg = "Model is still loading, please wait..." if model_manager.loading else (model_manager.load_error or "Model not loaded.")
        async def err():
            yield f"event: error\ndata: {json.dumps(msg)}\n\n"
        return StreamingResponse(err(), media_type="text/event-stream")

    session = _get_or_create_session(request.session_id, request.system_prompt)
    extract_user_facts(session, request.message)
    session.add_message("user", request.message)
    history = session.get_history_for_model()

    async def generate():
        full_response = []
        yield f"event: session\ndata: {json.dumps(session.session_id)}\n\n"

        for sse_data in model_manager.generate_stream(
            messages=history,
            temperature=request.temperature,
            max_tokens=request.max_tokens,
            enable_thinking=request.enable_thinking,
        ):
            yield sse_data
            try:
                for line in sse_data.strip().split("\n"):
                    if line.startswith("data: "):
                        token = json.loads(line[6:])
                        if isinstance(token, str):
                            full_response.append(token)
            except Exception:
                pass

        complete = "".join(full_response)
        if complete:
            # Clean response to remove any Gemma mentions
            complete = clean_response(complete)
            session.add_message("assistant", complete)

            # ── TRIGGER BACKGROUND COMPRESSION (ALL PAST) ──────
            _trigger_compression_async(session)

            # ── SUGGEST NEW CHAT IF NEEDED ──────────────────────
            if session.should_suggest_new_chat():
                yield f"event: suggest_new_chat\ndata: {json.dumps(True)}\n\n"

    return StreamingResponse(
        generate(),
        media_type="text/event-stream",
        headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"},
    )


# ── Sessions ─────────────────────────────────────────────────────
@app.get("/api/sessions")
async def list_sessions():
    return {"sessions": chat_engine.list_sessions()}


@app.get("/api/sessions/{session_id}")
async def get_session(session_id: str):
    session = chat_engine.get_session(session_id)
    if not session:
        raise HTTPException(status_code=404, detail="Session not found")
    return {
        "session_id": session_id,
        "messages": [m.to_dict() for m in session.messages],
        "summary": session.get_summary(),
    }


@app.delete("/api/sessions/{session_id}")
async def delete_session(session_id: str):
    if not chat_engine.delete_session(session_id):
        raise HTTPException(status_code=404, detail="Session not found")
    return {"success": True}


@app.post("/api/sessions/{session_id}/clear")
async def clear_session(session_id: str):
    session = chat_engine.get_session(session_id)
    if not session:
        raise HTTPException(status_code=404, detail="Session not found")
    session.clear_history()
    return {"success": True}


# ── Settings ─────────────────────────────────────────────────────
@app.get("/api/settings")
async def get_settings():
    return {"settings": {k: v for k, v in model_manager.settings.items() if k not in ["stop"]}}


@app.post("/api/settings")
async def update_settings(s: SettingsUpdate):
    if s.temperature is not None:
        model_manager.settings["temperature"] = s.temperature
    if s.top_p is not None:
        model_manager.settings["top_p"] = s.top_p
    if s.top_k is not None:
        model_manager.settings["top_k"] = s.top_k
    if s.max_tokens is not None:
        model_manager.settings["max_tokens"] = s.max_tokens
    return {"success": True, "settings": model_manager.settings}


# ── File Upload / Analysis ─────────────────────────────────────
@app.post("/api/analyze-file")
async def analyze_file(
    file: UploadFile = File(...),
    session_id: Optional[str] = Form(None),
    prompt: Optional[str] = Form(None)
):
    """
    Upload and analyze a file (image, PDF, TXT, DOC, DOCX).
    """
    print(f"\n{'='*50}")
    print(f"[UPLOAD] File: {file.filename}")
    print(f"[UPLOAD] Content-Type header: {file.content_type}")
    
    try:
        content = await file.read()
        print(f"[UPLOAD] Size: {len(content)} bytes")
    except Exception as e:
        print(f"[UPLOAD ERROR] Read failed: {e}")
        raise HTTPException(status_code=400, detail=f"Failed to read file: {e}")

    if len(content) == 0:
        raise HTTPException(status_code=400, detail="File is empty")

    # Get extension from filename (RELIABLE - not from MIME type)
    ext = file.filename.lower().split('.')[-1] if '.' in file.filename else ''
    print(f"[UPLOAD] Extension: '{ext}'")
    
    # Map extensions to types
    DOCUMENT_EXTS = ['txt', 'pdf', 'doc', 'docx']
    IMAGE_EXTS = ['png', 'jpg', 'jpeg', 'gif', 'bmp', 'webp', 'svg']
    
    is_document = ext in DOCUMENT_EXTS
    is_image = ext in IMAGE_EXTS
    
    print(f"[UPLOAD] is_document={is_document}, is_image={is_image}")

    if not is_document and not is_image:
        print(f"[UPLOAD ERROR] Unsupported extension: {ext}")
        raise HTTPException(status_code=400, detail=f"Unsupported file type: '{ext}'. Supported: {DOCUMENT_EXTS + IMAGE_EXTS}")

    # Save file to disk for persistence
    safe_filename = f"{uuid.uuid4().hex}_{file.filename}"
    file_path = uploads_path / safe_filename
    with open(file_path, "wb") as f:
        f.write(content)
    persistent_url = f"/uploads/{safe_filename}"
    print(f"[UPLOAD] Saved to {file_path}")

    # Create/get session
    session = _get_or_create_session(session_id, None)
    print(f"[UPLOAD] Session: {session.session_id}")

    try:
        from backend.chat_engine import process_upload, build_prompt
        
        file_text, is_image, needs_vision = process_upload(content, file.filename)
        
        if needs_vision:
            print("[UPLOAD] Routing image to Moondream2 vision model...")
            from backend.vision_processor import describe_image
            user_message = prompt or "What is in this image?"
            vision_answer = describe_image(content, user_message)
            
            response_text = sanitize_for_console(vision_answer)
            
            display_content = f"![{file.filename}]({persistent_url})\n\n{user_message}"
            session.add_message("user", sanitize_for_console(user_message), metadata={"display_content": display_content})
            session.add_message("assistant", response_text)
            print(f"{'='*50}\n")
            
            return {
                "response": response_text,
                "analysis": {"format": ext, "type": "vision", "analysis": "Processed by Moondream2"},
                "session_id": session.session_id,
                "ocr_preview": None
            }
        else:
            # Optional: truncate very long OCR output for 2B model context window
            if len(file_text) > 4000:  # safer limit after prompt wrapping overhead
                file_text = file_text[:4000] + "\n...[truncated]"
                
            user_message = prompt or "Please analyze this file."
            prompt_text = build_prompt(file_text, user_message, is_image)
            analysis_data = {
                "format": ext,
                "type": "processed",
                "analysis": f"Extracted {len(file_text)} characters"
            }

            # Sanitize and send to model
            prompt_text = sanitize_for_console(prompt_text)
            print(f"[UPLOAD] Prompt length: {len(prompt_text)} chars")

            if is_image:
                display_content = f"![{file.filename}]({persistent_url})\n\n{user_message}"
            else:
                display_content = f"📄 **{file.filename}**\n\n{user_message}" if user_message != "Please analyze this file." else f"📄 **{file.filename}**"
                
            session.add_message("user", prompt_text, metadata={"display_content": display_content})
            history = session.get_history_for_model()

        if not model_manager.model_loaded:
            raise HTTPException(status_code=503, detail="Model not loaded yet")

        print(f"[UPLOAD] Generating response...")
        result = model_manager.generate_sync(
            messages=history,
            temperature=1.0,
            max_tokens=2048,
        )
        
        if "error" in result:
            print(f"[UPLOAD ERROR] Model: {result['error']}")
            raise HTTPException(status_code=500, detail=f"Model error: {result['error']}")

        response_text = sanitize_for_console(result["response"])
        print(f"[UPLOAD] Response: {len(response_text)} chars")
        
        session.add_message("assistant", response_text)
        print(f"{'='*50}\n")

        return {
            "response": response_text,
            "analysis": analysis_data,
            "session_id": session.session_id,
            "ocr_preview": file_text[:500] if file_text else None
        }
        
    except HTTPException:
        raise
    except Exception as e:
        import traceback
        print(f"[UPLOAD ERROR] {e}")
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Processing error: {str(e)}")


# Alias endpoint for backward compatibility with cached frontend
@app.post("/api/analyze-image")
async def analyze_image(file: UploadFile = File(...), session_id: Optional[str] = Form(None), prompt: Optional[str] = Form(None)):
    return await analyze_file(file, session_id, prompt)


# ── Helper functions ───────────────────────────────────────────
def _trigger_compression_async(session: ChatSession):
    """Trigger history compression in background thread after response."""
    def compress_task():
        try:
            from backend.model_manager import model_manager
            if model_manager.model_loaded and not model_manager.loading:
                did_compress = session.check_and_compress(model_manager)
                if did_compress:
                    print(f"[COMPRESSION] Session {session.session_id}: compressed successfully")
        except Exception as e:
            print(f"[COMPRESSION BG ERROR] {e}")
    
    thread = threading.Thread(target=compress_task, daemon=True)
    thread.start()


def _get_or_create_session(session_id, system_prompt) -> ChatSession:
    if session_id:
        session = chat_engine.get_session(session_id)
        if session:
            # Check if the frontend's raw system prompt changed
            # (We check against the frontend's string if possible, or just force update)
            if system_prompt:
                session.set_system_prompt(system_prompt)
            return session
    return chat_engine.create_session(session_id, system_prompt)


def extract_user_facts(session: ChatSession, message: str):
    """Detect and store user facts from their message."""
    name_match = re.search(
        r"(?:my name is|i am|i'm|call me)\s+([A-Za-z][a-z]+(?:\s+[A-Za-z][a-z]+)?)",
        message,
        re.IGNORECASE,
    )
    if name_match:
        name = name_match.group(1).strip()
        session.update_user_fact("User's name", name)
        print(f"[MEMORY] Stored user name: {name}")


def clean_response(response_text: str) -> str:
    """
    Post-process response to remove any Gemma/Google model references.
    This handles cases where the model tries to override its identity.
    This is a final safety net after token-level cleaning.
    """
    # Replace model names
    response_text = re.sub(r'\bGemma\s*\d*\.?\d*\b', 'Genie', response_text, flags=re.IGNORECASE)
    response_text = re.sub(r'\bgemini\b', 'Genie', response_text, flags=re.IGNORECASE)
    response_text = re.sub(r'\bGoogle\b', 'Genie', response_text, flags=re.IGNORECASE)
    
    # Remove Google training references
    response_text = re.sub(r',?\s*trained by[^.]*?\.', '.', response_text, flags=re.IGNORECASE)
    response_text = re.sub(r',?\s*developed by[^.]*?\.', '.', response_text, flags=re.IGNORECASE)
    response_text = re.sub(r',?\s*created by[^.]*?\.', '.', response_text, flags=re.IGNORECASE)
    
    # Replace or remove "don't have a name" phrases
    response_text = re.sub(r"[,\s]+I\s+don't\s+have\s+a\s+(personal\s+)?name", '', response_text, flags=re.IGNORECASE)
    response_text = re.sub(r"[,\s]+I\s+don't\s+have\s+a\s+name", '', response_text, flags=re.IGNORECASE)
    
    # Replace LLM identity statements
    response_text = re.sub(r'I\s+am\s+a\s+large\s+language\s+model', 'I am Genie', response_text, flags=re.IGNORECASE)
    response_text = re.sub(r'I\s+am\s+(an?\s+)?language\s+model', 'I am Genie', response_text, flags=re.IGNORECASE)
    response_text = re.sub(r'I\s+am\s+(an?\s+)?(AI|artificial intelligence)\s+model', 'I am Genie', response_text, flags=re.IGNORECASE)
    response_text = re.sub(r"I\'m\s+a\s+(?:large\s+)?language\s+model", "I'm Genie", response_text, flags=re.IGNORECASE)
    response_text = re.sub(r'as\s+a\s+(?:large\s+)?language\s+model', 'as Genie', response_text, flags=re.IGNORECASE)
    
    # Remove hallucinated control tokens like <end_of_turn> or <channel|>
    response_text = re.sub(r'<[^>]+(?:_turn|\|)>', '', response_text, flags=re.IGNORECASE)
    
    return response_text


def analyze_image_bytes(data: bytes) -> dict:
    fmt = 'unknown'
    width = height = None
    mode = None
    aspect_ratio = None

    if Image is not None:
        try:
            with Image.open(io.BytesIO(data)) as img:
                fmt = (img.format or 'unknown').lower()
                width, height = img.size
                mode = img.mode
                if width and height:
                    aspect_ratio = round(width / height, 2)
        except Exception:
            width = height = None

    size_kb = round(len(data) / 1024, 2)
    analysis = f"Detected {fmt.upper()} image"
    if width and height:
        analysis += f" with dimensions {width}×{height}"
    if mode:
        analysis += f" in {mode} mode"
    if aspect_ratio:
        analysis += f" (aspect ratio {aspect_ratio}:1)"
    analysis += f" and size {size_kb} KB."
    if size_kb > 512:
        analysis += " This image is large and may take longer to process."
    elif size_kb < 100:
        analysis += " This is a small image."

    return {
        "format": fmt,
        "width": width,
        "height": height,
        "mode": mode,
        "aspect_ratio": aspect_ratio,
        "size_kb": size_kb,
        "analysis": analysis,
    }


def clean_extracted_text(text: str) -> str:
    """Clean up extracted text by removing excessive whitespace and artifacts."""
    import re
    
    # Remove excessive blank lines (more than 2)
    text = re.sub(r'\n{4,}', '\n\n\n', text)
    
    # Normalize spaces (but preserve intentional indentation)
    lines = text.split('\n')
    cleaned_lines = []
    for line in lines:
        # Strip trailing spaces, keep leading for structure
        cleaned = line.rstrip()
        cleaned_lines.append(cleaned)
    
    text = '\n'.join(cleaned_lines)
    
    # Remove null bytes and weird characters
    text = text.replace('\x00', '').replace('\x0b', '').replace('\x0c', '')
    
    return text.strip()


def extract_document_text(filename: str, data: bytes) -> str:
    """
    Extract text from PDF, TXT, DOC, DOCX files.
    Returns extracted text or error message in [brackets].
    """
    ext = filename.lower().split('.')[-1] if '.' in filename else ''
    print(f"[DOC] Processing .{ext} file ({len(data)} bytes)")
    
    text = ""

    try:
        # ── TXT files ──────────────────────────────────────────
        if ext == 'txt':
            print("[DOC] Decoding TXT file...")
            
            # Try multiple encodings in order of likelihood
            encodings = ['utf-8', 'utf-8-sig', 'latin-1', 'cp1252', 'iso-8859-1', 'utf-16']
            
            for encoding in encodings:
                try:
                    text = data.decode(encoding, errors='strict')
                    print(f"[DOC] Decoded with {encoding}: {len(text)} chars")
                    break
                except (UnicodeDecodeError, UnicodeError):
                    continue
            
            # Fallback: force decode with replacement
            if not text:
                text = data.decode('utf-8', errors='replace')
                print(f"[DOC] Force decoded (with replacements): {len(text)} chars")
                
            # Remove null bytes and control chars
            text = text.replace('\x00', '').replace('\x0b', '').replace('\x0c', '')
            
        # ── PDF files ────────────────────────────────────────────
        elif ext == 'pdf':
            if PyPDF2 is None:
                return "[ERROR: PyPDF2 not installed. Run: pip install PyPDF2==3.0.1]"
            
            try:
                reader = PyPDF2.PdfReader(io.BytesIO(data))
                total_pages = len(reader.pages)
                print(f"[DOC] PDF has {total_pages} pages")
                
                # Process first 3 pages only (speed optimization)
                pages_to_process = min(3, total_pages)
                for i in range(pages_to_process):
                    try:
                        page = reader.pages[i]
                        page_text = page.extract_text()
                        if page_text:
                            text += f"\n--- Page {i+1} ---\n{page_text}\n"
                    except Exception as e:
                        print(f"[DOC] Page {i} failed: {e}")
                        continue
                
                if total_pages > pages_to_process:
                    text += f"\n[... {total_pages - pages_to_process} more pages not read ...]"
                    
                print(f"[DOC] PDF extracted: {len(text)} chars")
                
            except Exception as e:
                return f"[ERROR reading PDF: {e}]"

        # ── DOC/DOCX files ───────────────────────────────────────
        elif ext in ['doc', 'docx']:
            if ext == 'docx':
                if docx is None:
                    return "[ERROR: python-docx not installed. Run: pip install python-docx==1.1.2]"
                
                try:
                    document = docx.Document(io.BytesIO(data))
                    
                    # Extract text from all paragraphs
                    paragraphs = []
                    for para in document.paragraphs:
                        if para.text.strip():
                            paragraphs.append(para.text)
                    
                    text = '\n'.join(paragraphs)
                    print(f"[DOC] DOCX extracted: {len(text)} chars from {len(paragraphs)} paragraphs")
                    
                    # Extract from headers and footers
                    try:
                        for section in document.sections:
                            header = section.header
                            if header:
                                for para in header.paragraphs:
                                    if para.text.strip():
                                        text += f"\n[HEADER] {para.text}"
                            
                            footer = section.footer
                            if footer:
                                for para in footer.paragraphs:
                                    if para.text.strip():
                                        text += f"\n[FOOTER] {para.text}"
                    except Exception as hf_error:
                        print(f"[DOC] Header/footer extraction failed: {hf_error}")
                    
                    # Also try to extract from tables
                    table_text = []
                    for table in document.tables:
                        for row in table.rows:
                            row_text = ' | '.join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                            if row_text:
                                table_text.append(row_text)
                    
                    if table_text:
                        text += '\n\n[TABLES]\n' + '\n'.join(table_text)
                        print(f"[DOC] Added {len(table_text)} table rows")
                    
                    # If still no text, try to get all text using different methods
                    if len(text.strip()) < 10:
                        print("[DOC] Trying alternative extraction methods...")
                        try:
                            alt_text = extract_docx_text_from_zip(data)
                            if alt_text and len(alt_text.strip()) > len(text):
                                text = alt_text
                                print(f"[DOC] Zip extraction found {len(text)} chars")
                        except Exception as alt_error:
                            print(f"[DOC] Alternative extraction failed: {alt_error}")

                    if len(text.strip()) == 0:
                        return "[ERROR: No readable text found in DOCX file. The document may contain only images, be password-protected, or be corrupted.]"
                        
                except Exception as e:
                    return f"[ERROR reading DOCX: {e}]"
            
            elif ext == 'doc':
                # Handle older .doc files
                if olefile is None:
                    return "[ERROR: olefile not installed. Run: pip install olefile]"
                
                try:
                    # Try to extract text from .doc file using olefile
                    ole = olefile.OleFileIO(io.BytesIO(data))
                    
                    # Look for WordDocument stream
                    if ole.exists('WordDocument'):
                        stream = ole.openstream('WordDocument')
                        doc_data = stream.read()
                        
                        # Basic text extraction from .doc (very limited)
                        # This is a simplified approach - real .doc parsing is complex
                        try:
                            # Try to decode as UTF-16 (common in .doc files)
                            text = doc_data.decode('utf-16', errors='ignore')
                            # Clean up control characters
                            text = ''.join(c for c in text if ord(c) >= 32 or c in '\n\r\t')
                            text = text.strip()
                            
                            if text:
                                print(f"[DOC] DOC extracted: {len(text)} chars")
                            else:
                                return "[ERROR: No readable text found in DOC file. The document may be corrupted or use an unsupported format.]"
                        except:
                            return "[ERROR: Cannot decode DOC file content. Try saving as DOCX format.]"
                    else:
                        return "[ERROR: Invalid DOC file format - missing WordDocument stream.]"
                        
                except Exception as e:
                    return f"[ERROR reading DOC file: {e}]"
        
        else:
            return f"[ERROR: Unsupported extension '{ext}']"

    except Exception as e:
        return f"[ERROR extracting text: {e}]"

    # Clean up extracted text
    text = clean_extracted_text(text)
    
    # Truncate if too long for context window
    MAX_CHARS = 6000  # Safe for 8K context
    if len(text) > MAX_CHARS:
        remaining = len(text) - MAX_CHARS
        text = text[:MAX_CHARS] + f"\n\n[... {remaining} more characters truncated ...]"
        print(f"[DOC] Truncated to {MAX_CHARS} chars")

    print(f"[DOC] Final output: {len(text)} characters")
    return text.strip()


def extract_docx_text_from_zip(data: bytes) -> str:
    """Fallback extraction for DOCX text using ZIP/XML parsing."""
    try:
        with zipfile.ZipFile(io.BytesIO(data)) as z:
            candidate_files = [
                'word/document.xml',
                'word/footnotes.xml',
                'word/endnotes.xml',
                'word/comments.xml',
            ]
            candidate_files += [name for name in z.namelist() if name.startswith('word/header') or name.startswith('word/footer')]
            seen = set()
            text_fragments = []
            for name in candidate_files:
                if name in seen:
                    continue
                seen.add(name)
                try:
                    xml_bytes = z.read(name)
                except KeyError:
                    continue
                try:
                    root = ET.fromstring(xml_bytes)
                except ET.ParseError:
                    continue
                for element in root.iter():
                    tag = element.tag
                    if tag.endswith('}t') and element.text:
                        text_fragments.append(element.text)
                    elif tag.endswith('}tab'):
                        text_fragments.append('\t')
                    elif tag.endswith('}br') or tag.endswith('}cr'):
                        text_fragments.append('\n')
                text_fragments.append('\n')
            return ''.join(text_fragments).strip()
    except Exception as e:
        print(f"[DOC] ZIP fallback failed: {e}")
        return ''


def build_image_prompt(file: UploadFile, analysis: dict) -> str:
    prompt_lines = [
        f"The user uploaded an image named '{file.filename}'.",
        f"Image type: {analysis['format'].upper()}",
    ]

    if analysis.get('width') and analysis.get('height'):
        prompt_lines.append(f"Image dimensions: {analysis['width']}×{analysis['height']}")
    if analysis.get('mode'):
        prompt_lines.append(f"Image mode: {analysis['mode']}")
    if analysis.get('aspect_ratio') is not None:
        prompt_lines.append(f"Aspect ratio: {analysis['aspect_ratio']}:1")
    prompt_lines.append(f"Image size: {analysis['size_kb']} KB.")
    prompt_lines.append(
        "Please analyze the image content and describe what you can infer about the objects, scene, and layout."
    )
    prompt_lines.append(
        "Use these image details to provide a clear, visual analysis as if you had seen the image directly."
    )
    return "\n".join(prompt_lines)


if __name__ == "__main__":
    import uvicorn
    print("=" * 60)
    print("  Genie Chatbot Server")
    print("  API: http://localhost:8000")
    print("  Frontend: http://localhost:8000/app")
    print("=" * 60)
    uvicorn.run("backend.main:app", host="0.0.0.0", port=8000, reload=False, log_level="info")